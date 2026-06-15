#!/usr/bin/env python3
"""Generate docs/系统技术文档.docx from the canonical Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "系统技术文档.md"
OUTPUT = ROOT / "docs" / "系统技术文档.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x1F, 0x38, 0x64)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "E8EEF5"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT_LATIN = "Arial"
FONT_EAST_ASIA = "PingFang SC"


def set_run_font(run, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = FONT_LATIN
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (28, NAVY, 0, 8),
        "Subtitle": (13, MUTED, 0, 10),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("大宗物资供应链全链路管理系统  |  技术文档")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def create_numbering_sequence(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level_type)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("w:start", "1"),
        ("w:numFmt", "decimal"),
        ("w:lvlText", "%1."),
        ("w:lvlJc", "left"),
    ):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        level.append(node)
    paragraph_properties = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)
    abstract.append(level)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    existing = paragraph_properties.find(qn("w:numPr"))
    if existing is not None:
        paragraph_properties.remove(existing)
    numbering_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    numbering_properties.append(level)
    numbering_properties.append(number)
    paragraph_properties.append(numbering_properties)


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=9.5, color=DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    base = TABLE_WIDTH_DXA // cols
    widths = [base] * cols
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_inline_markdown(p, row[c_idx] if c_idx < len(row) else "")
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for run in p.runs:
                    run.bold = True
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def is_separator_row(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        if not is_separator_row(lines[idx]):
            rows.append([cell.strip() for cell in lines[idx].strip().strip("|").split("|")])
        idx += 1
    return rows, idx


def build_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = next(line[2:].strip() for line in lines if line.startswith("# "))
    section_names = [line[3:].strip() for line in lines if line.startswith("## ")]

    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_markdown(p, title)
    p = doc.add_paragraph("系统架构、业务流程、状态模型、数据领域、接口、集成与运维", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_markdown(meta, "版本 3.0  |  代码核验 2026-06-12  |  Markdown 单一事实源")
    doc.add_page_break()

    doc.add_heading("目录", level=1)
    toc_num_id = create_numbering_sequence(doc)
    for name in section_names:
        p = doc.add_paragraph(style="List Number")
        apply_numbering(p, toc_num_id)
        add_inline_markdown(p, re.sub(r"^\d+\.\s*", "", name))
    doc.add_page_break()

    idx = 0
    in_code = False
    code_lines: list[str] = []
    skipped_title = False
    numbered_list_id: int | None = None
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.15)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_together = True
                for pos, code_line in enumerate(code_lines):
                    run = p.add_run(code_line + ("\n" if pos < len(code_lines) - 1 else ""))
                    set_run_font(run, size=9, color=DARK_BLUE)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if line.startswith("# ") and not skipped_title:
            skipped_title = True
            idx += 1
            continue
        if stripped.startswith(">"):
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            text = heading.group(2)
            p = doc.add_heading(level=level)
            add_inline_markdown(p, text)
            idx += 1
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, bullet.group(1))
            idx += 1
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            if numbered_list_id is None:
                numbered_list_id = create_numbering_sequence(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, numbered_list_id)
            add_inline_markdown(p, numbered.group(1))
            idx += 1
            continue
        numbered_list_id = None
        if stripped:
            p = doc.add_paragraph()
            add_inline_markdown(p, stripped)
        idx += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT} from {SOURCE}")


if __name__ == "__main__":
    build_document()
